"""Tests for the CPU-side components (features, classifier, evaluate).

Uses synthetic corpora that mimic the key stylometric contrast:
"AI-like" text with uniform sentence lengths and low burstiness vs
"human-like" text with high sentence-length variance.
"""

import random

import numpy as np
import pytest

from aidetect import (
    FEATURE_NAMES,
    FeatureDetector,
    evaluate,
    extract_features,
    pick_threshold,
)

random.seed(42)

VOCAB = (
    "system data model process value result method network analysis point "
    "structure element function change level time work part case number "
    "quickly slowly carefully suddenly the a of to in on with and or but"
).split()


def make_sentence(n_words: int) -> str:
    return " ".join(random.choices(VOCAB, k=n_words)).capitalize() + "."


def make_ai_like() -> str:
    # Uniform sentence lengths -> low burstiness
    return " ".join(make_sentence(random.randint(14, 16)) for _ in range(8))


def make_human_like() -> str:
    # Highly variable sentence lengths -> high burstiness
    return " ".join(make_sentence(random.choice([3, 5, 8, 22, 30])) for _ in range(8))


@pytest.fixture(scope="module")
def corpus():
    ai = [make_ai_like() for _ in range(120)]
    human = [make_human_like() for _ in range(120)]
    texts = ai + human
    labels = [1] * len(ai) + [0] * len(human)
    idx = list(range(len(texts)))
    random.shuffle(idx)
    texts = [texts[i] for i in idx]
    labels = [labels[i] for i in idx]
    split = int(0.7 * len(texts))
    return (texts[:split], labels[:split]), (texts[split:], labels[split:])


def test_feature_vector_shape_and_finiteness():
    vec = extract_features("Hello world. This is a test sentence, with a comma!")
    assert len(vec) == len(FEATURE_NAMES)
    assert all(np.isfinite(v) for v in vec)


def test_features_handle_edge_cases():
    for text in ["", "word", "!!!", "Μία ελληνική πρόταση εδώ."]:
        vec = extract_features(text)
        assert len(vec) == len(FEATURE_NAMES)
        assert all(np.isfinite(v) for v in vec)


def test_burstiness_separates_corpora():
    b_idx = FEATURE_NAMES.index("burstiness")
    ai_b = np.mean([extract_features(make_ai_like())[b_idx] for _ in range(30)])
    hu_b = np.mean([extract_features(make_human_like())[b_idx] for _ in range(30)])
    assert hu_b > ai_b


def test_classifier_trains_and_separates(corpus):
    (Xtr, ytr), (Xte, yte) = corpus
    det = FeatureDetector().fit(Xtr, ytr)
    probs = det.predict_proba(Xte)
    report = evaluate(probs, yte, max_fpr=0.05)
    assert report.auroc > 0.9, f"AUROC too low: {report}"


def test_save_load_roundtrip(tmp_path, corpus):
    (Xtr, ytr), (Xte, _) = corpus
    det = FeatureDetector().fit(Xtr, ytr)
    p = tmp_path / "model.pkl"
    det.save(p)
    det2 = FeatureDetector.load(p)
    np.testing.assert_allclose(det.predict_proba(Xte), det2.predict_proba(Xte))


def test_pick_threshold_respects_fpr(corpus):
    (Xtr, ytr), (Xte, yte) = corpus
    det = FeatureDetector().fit(Xtr, ytr)
    probs = det.predict_proba(Xte)
    thr = pick_threshold(probs, np.array(yte), max_fpr=0.05)
    report = evaluate(probs, yte, threshold=thr)
    assert report.fpr_at_threshold <= 0.05 + 1e-9


def test_evaluate_perfect_separation():
    scores = [0.9, 0.8, 0.95, 0.1, 0.2, 0.05]
    labels = [1, 1, 1, 0, 0, 0]
    report = evaluate(scores, labels)
    assert report.auroc == 1.0
    assert report.fpr_at_threshold == 0.0
    assert report.tpr_at_threshold == 1.0


# --- ensemble & max-accuracy threshold tests ---

from aidetect.ensemble import ScoreEnsemble, chunk_text, score_long_text
from aidetect.evaluate import pick_threshold_max_accuracy


def _two_detector_scores(n=200, seed=7):
    rng = np.random.default_rng(seed)
    labels = rng.integers(0, 2, n)
    # detector A: decent, detector B: decent but errors are independent
    a = labels + rng.normal(0, 0.8, n)
    b = labels + rng.normal(0, 0.8, n)
    return np.column_stack([a, b]), labels


def test_ensemble_beats_single_detector():
    X, y = _two_detector_scores()
    Xtr, ytr, Xte, yte = X[:140], y[:140], X[140:], y[140:]
    ens = ScoreEnsemble().fit(Xtr, ytr)
    from sklearn.metrics import roc_auc_score
    auc_a = roc_auc_score(yte, Xte[:, 0])
    auc_ens = roc_auc_score(yte, ens.predict_proba(Xte))
    assert auc_ens > auc_a


def test_ensemble_save_load(tmp_path):
    X, y = _two_detector_scores()
    ens = ScoreEnsemble().fit(X, y)
    p = tmp_path / "ens.pkl"
    ens.save(p)
    np.testing.assert_allclose(
        ens.predict_proba(X), ScoreEnsemble.load(p).predict_proba(X)
    )


def test_max_accuracy_threshold_beats_fpr_constrained_on_accuracy():
    rng = np.random.default_rng(3)
    labels = rng.integers(0, 2, 400)
    scores = labels + rng.normal(0, 0.7, 400)
    thr_acc = pick_threshold_max_accuracy(scores, labels)
    thr_fpr = pick_threshold(scores, np.asarray(labels), max_fpr=0.01)
    acc = lambda t: ((scores >= t).astype(int) == labels).mean()
    assert acc(thr_acc) >= acc(thr_fpr)


def test_chunking_and_aggregation():
    long_text = " ".join(["word"] * 1000)
    chunks = chunk_text(long_text, chunk_words=300, overlap=50)
    assert len(chunks) > 1
    result = score_long_text(lambda c: float(len(c.split())), long_text)
    assert result["n_chunks"] == len(chunks)
    assert result["max"] >= result["mean"]


# --- document reader tests ---

from aidetect.readers import load_directory, load_text


def _make_docx(path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph of the body text.")
    doc.add_paragraph("Second paragraph, with a comma and more words.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    doc.save(str(path))


def test_read_docx_paragraphs_and_tables(tmp_path):
    p = tmp_path / "sample.docx"
    _make_docx(p)
    text = load_text(p)
    assert "First paragraph" in text
    assert "Second paragraph" in text
    assert "Cell A | Cell B" in text


def test_docx_feeds_into_feature_pipeline(tmp_path):
    p = tmp_path / "sample.docx"
    _make_docx(p)
    vec = extract_features(load_text(p))
    assert len(vec) == len(FEATURE_NAMES)
    assert all(np.isfinite(v) for v in vec)


def test_load_text_txt_and_unsupported(tmp_path):
    t = tmp_path / "a.txt"
    t.write_text("plain text here", encoding="utf-8")
    assert load_text(t) == "plain text here"
    with pytest.raises(ValueError, match="Unsupported"):
        load_text(tmp_path / "a.xyz")
    with pytest.raises(ValueError, match="Legacy .doc"):
        load_text(tmp_path / "old.doc")


def test_load_directory(tmp_path):
    (tmp_path / "a.txt").write_text("alpha", encoding="utf-8")
    _make_docx(tmp_path / "b.docx")
    (tmp_path / "ignore.xyz").write_text("nope", encoding="utf-8")
    files = load_directory(tmp_path)
    assert set(files) == {"a.txt", "b.docx"}
