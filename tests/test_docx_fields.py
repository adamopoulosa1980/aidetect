"""Tests for Word field filtering and Markdown export.

Word generates field results the author never typed — table-of-contents
entries, cross-references, and placeholders like "No table of figures entries
found." Feeding those to a detector dilutes the signal, so they are stripped.
They are matched structurally rather than by their text, because Word localises
the wording.
"""

from __future__ import annotations

import pytest

docx = pytest.importorskip("docx")

from aidetect.readers import docx_to_markdown, read_docx, save_markdown  # noqa: E402


def _add_field(paragraph, instr: str, result: str) -> None:
    """Append a Word field in its usual on-disk form.

    fldChar begin / instrText / fldChar separate / result run / fldChar end.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _run(child):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    _run(begin)

    instruction = OxmlElement("w:instrText")
    instruction.text = instr
    _run(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    _run(separate)

    text = OxmlElement("w:t")
    text.text = result
    _run(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _run(end)


PROSE = "The study measured rainfall at three upland sites each season."


@pytest.fixture
def document_with_fields(tmp_path):
    d = docx.Document()
    _add_field(d.add_paragraph(), r'TOC \h \z \c "Figure"', "No table of figures entries found.")
    _add_field(d.add_paragraph(), "PAGEREF _Toc123 \\h", "Error! Reference source not found.")
    d.add_paragraph(PROSE)
    path = tmp_path / "fields.docx"
    d.save(path)
    return path


def test_field_placeholders_are_stripped(document_with_fields):
    text = read_docx(document_with_fields)
    assert "No table of figures entries found." not in text
    assert "Error! Reference source not found." not in text
    assert PROSE in text


def test_localised_placeholder_is_stripped_too(tmp_path):
    """The Greek wording must go as well, which text matching would miss."""
    d = docx.Document()
    _add_field(d.add_paragraph(), 'TOC \\c "Figure"', "Δεν βρέθηκαν καταχωρήσεις.")
    d.add_paragraph(PROSE)
    path = tmp_path / "greek.docx"
    d.save(path)

    text = read_docx(path)
    assert "Δεν βρέθηκαν" not in text
    assert PROSE in text


def test_toc_styled_paragraphs_are_dropped(tmp_path):
    d = docx.Document()
    para = d.add_paragraph("1. Introduction\t4")
    try:
        para.style = d.styles["TOC 1"]
    except KeyError:
        pytest.skip("template has no TOC style")
    d.add_paragraph(PROSE)
    path = tmp_path / "toc.docx"
    d.save(path)

    text = read_docx(path)
    assert "Introduction" not in text
    assert PROSE in text


def test_plain_documents_are_unaffected(tmp_path):
    d = docx.Document()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    path = tmp_path / "plain.docx"
    d.save(path)
    assert read_docx(path) == "First paragraph.\nSecond paragraph."


# ------------------------------------------------------------------ markdown


def test_markdown_keeps_heading_structure(tmp_path):
    d = docx.Document()
    d.add_heading("Report Title", 0)
    d.add_heading("Section One", 1)
    d.add_paragraph(PROSE)
    path = tmp_path / "structured.docx"
    d.save(path)

    md = docx_to_markdown(path)
    assert "# Report Title" in md
    assert "## Section One" in md
    assert PROSE in md


def test_markdown_renders_tables(tmp_path):
    d = docx.Document()
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "AUROC"
    table.rows[1].cells[1].text = "0.94"
    path = tmp_path / "table.docx"
    d.save(path)

    md = docx_to_markdown(path)
    assert "| Metric | Value |" in md
    assert "| --- | --- |" in md
    assert "| AUROC | 0.94 |" in md


def test_markdown_export_also_drops_fields(document_with_fields):
    md = docx_to_markdown(document_with_fields)
    assert "No table of figures entries found." not in md
    assert PROSE in md


def test_save_markdown_writes_beside_the_source(document_with_fields):
    written = save_markdown(document_with_fields)
    assert written.suffix == ".md"
    assert written.parent == document_with_fields.parent
    assert PROSE in written.read_text(encoding="utf-8")


def test_save_markdown_refuses_to_overwrite_its_source(tmp_path):
    source = tmp_path / "notes.md"
    source.write_text("# Notes\n", encoding="utf-8")
    with pytest.raises(ValueError, match="Refusing to overwrite"):
        save_markdown(source)
